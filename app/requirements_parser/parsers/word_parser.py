"""
Word文档解析器
解析Word格式的需求文档
"""
import re
from typing import List, Dict, Any, Optional
from datetime import datetime
from pathlib import Path

try:
    from docx import Document as DocxDocument
    import docx
except ImportError:
    DocxDocument = None
    docx = None

from app.requirements_parser.parsers.base import BaseParser
from app.requirements_parser.models.document import Document, DocumentType, DocumentMetadata


class WordParser(BaseParser):
    """Word文档解析器"""
    
    def __init__(self):
        """初始化Word解析器"""
        super().__init__()
        self.supported_extensions = {'.docx', '.doc'}
        
        # 检查依赖
        if docx is None:
            raise ImportError("需要安装python-docx库: pip install python-docx")
    
    def parse(self, content: str, **kwargs) -> Document:
        """
        解析Word内容（从文件路径）
        
        Args:
            content: Word文件路径
            **kwargs: 额外参数
            
        Returns:
            Document: 解析后的文档对象
        """
        if isinstance(content, str) and Path(content).exists():
            return self.parse_from_file(content, **kwargs)
        else:
            raise ValueError("Word解析器需要文件路径作为输入")
    
    def parse_from_file(self, file_path: str, encoding: str = "utf-8", **kwargs) -> Document:
        """
        从Word文件解析文档
        
        Args:
            file_path: Word文件路径
            encoding: 编码（Word不需要，保持接口一致性）
            **kwargs: 额外参数
            
        Returns:
            Document: 解析后的文档对象
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Word文件不存在: {file_path}")
        
        if not path.is_file():
            raise ValueError(f"路径不是文件: {file_path}")
        
        if path.suffix.lower() not in {'.docx', '.doc'}:
            raise ValueError(f"不是Word文件: {file_path}")
        
        # 提取Word文档内容
        paragraphs = self._extract_paragraphs_from_word(file_path)
        
        # 解析段落内容
        document = self._parse_paragraphs(paragraphs, str(path))
        document.file_path = str(path)
        
        # 设置文件元数据
        if document.metadata:
            document.metadata.file_size = path.stat().st_size
            document.metadata.created_at = path.stat().st_ctime
            document.metadata.modified_at = path.stat().st_mtime
            document.metadata.encoding = "binary"  # Word是二进制格式
        
        return document
    
    def _extract_paragraphs_from_word(self, file_path: str) -> List[str]:
        """
        从Word文件提取段落内容
        
        Args:
            file_path: Word文件路径
            
        Returns:
            List[str]: 段落文本列表
        """
        try:
            doc = DocxDocument(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                paragraphs.append(text)
            
            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        paragraphs.append(row_text)
            
            return paragraphs
            
        except Exception as e:
            raise ValueError(f"无法解析Word文件 {file_path}: {e}")
    
    def _parse_paragraphs(self, paragraphs: List[str], file_name: str) -> Document:
        """
        解析Word文档段落
        
        Args:
            paragraphs: 段落列表
            file_name: 文件名
            
        Returns:
            Document: 解析后的文档对象
        """
        # 验证内容
        if not paragraphs:
            raise ValueError("Word文档内容为空")
        
        # 过滤空段落
        non_empty_paragraphs = [p for p in paragraphs if p.strip()]
        if not non_empty_paragraphs:
            raise ValueError("Word文档内容为空")
        
        # 合并段落为文本内容
        content = '\n'.join(paragraphs)
        
        # 提取标题
        title = self._extract_title_from_paragraphs(non_empty_paragraphs, file_name)
        
        # 创建文档对象
        document = Document(
            title=title,
            content=content,
            document_type=DocumentType.WORD
        )
        
        # 添加解析的结构化信息
        document.sections = self._extract_sections_from_paragraphs(paragraphs)
        document.tables = self._extract_tables_from_paragraphs(paragraphs)
        document.links = self._extract_links_from_content(content)
        document.user_stories = self._extract_user_stories_from_content(content)
        
        return document
    
    def _extract_title_from_paragraphs(self, paragraphs: List[str], file_name: str) -> str:
        """
        从段落中提取标题
        
        Args:
            paragraphs: 非空段落列表
            file_name: 文件名
            
        Returns:
            str: 文档标题
        """
        if not paragraphs:
            return Path(file_name).stem
        
        # 第一个非空段落通常是标题
        first_paragraph = paragraphs[0].strip()
        
        # 如果第一段太长，可能不是标题
        if len(first_paragraph) > 100:
            # 查找更短的段落作为标题
            for paragraph in paragraphs[:5]:  # 只检查前5段
                if paragraph.strip() and len(paragraph.strip()) < 50:
                    return paragraph.strip()
        
        return first_paragraph if first_paragraph else Path(file_name).stem
    
    def _extract_sections_from_paragraphs(self, paragraphs: List[str]) -> List[Dict[str, Any]]:
        """
        从段落中提取章节信息
        
        Args:
            paragraphs: 段落列表
            
        Returns:
            List[Dict[str, Any]]: 章节信息列表
        """
        sections = []
        
        for i, paragraph in enumerate(paragraphs):
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            # 检测章节标题的模式
            section_patterns = [
                (r'^(\d+(?:\.\d+)*\.?)\s+(.+)$', 'numbered'),  # 1. 1.1 1.1.1 格式
                (r'^[一二三四五六七八九十]+[、．]\s*(.+)$', 'chinese'),  # 中文数字格式
                (r'^[A-Z]\.\s+(.+)$', 'alpha'),  # A. B. C. 格式
                (r'^第[一二三四五六七八九十]+章\s*(.*)$', 'chapter'),  # 第X章格式
            ]

            for pattern, pattern_type in section_patterns:
                match = re.match(pattern, paragraph)
                if match:
                    # 确定章节级别和标题
                    if pattern_type == 'chapter':
                        level = 1
                        title = match.group(1) if match.group(1) else paragraph
                    elif pattern_type == 'numbered':
                        # 计算数字层级 (1. = 1, 1.1 = 2, 1.1.1 = 3)
                        number_part = match.group(1).rstrip('.')
                        level = number_part.count('.') + 1
                        title = match.group(2)
                    elif pattern_type == 'chinese':
                        level = 2
                        title = match.group(1)
                    elif pattern_type == 'alpha':
                        level = 2
                        title = match.group(1)
                    else:
                        level = 2
                        title = match.group(1)

                    sections.append({
                        'title': title.strip(),
                        'level': level,
                        'line_number': i + 1
                    })
                    break
        
        return sections
    
    def _extract_tables_from_paragraphs(self, paragraphs: List[str]) -> List[Dict[str, Any]]:
        """
        从段落中提取表格信息
        
        Args:
            paragraphs: 段落列表
            
        Returns:
            List[Dict[str, Any]]: 表格信息列表
        """
        tables = []
        
        for i, paragraph in enumerate(paragraphs):
            # 检测表格行（包含|字符）
            if '|' in paragraph and paragraph.count('|') >= 2:
                tables.append({
                    'content': paragraph.strip(),
                    'line_number': i + 1
                })
        
        return tables
    
    def _extract_links_from_content(self, content: str) -> List[Dict[str, Any]]:
        """
        从内容中提取链接信息
        
        Args:
            content: 文档内容
            
        Returns:
            List[Dict[str, Any]]: 链接信息列表
        """
        links = []
        
        # 查找URL
        url_pattern = r'https?://[^\s]+'
        for match in re.finditer(url_pattern, content):
            links.append({
                'url': match.group(),
                'position': match.start()
            })
        
        # 查找邮箱
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        for match in re.finditer(email_pattern, content):
            links.append({
                'email': match.group(),
                'position': match.start()
            })
        
        return links
    
    def _extract_user_stories_from_content(self, content: str) -> List[Dict[str, Any]]:
        """
        从内容中提取用户故事
        
        Args:
            content: 文档内容
            
        Returns:
            List[Dict[str, Any]]: 用户故事列表
        """
        user_stories = []
        
        # 用户故事模式
        story_patterns = [
            r'作为.*?，我希望.*?，以便.*?[。.]',
            r'As a.*?, I want.*?, so that.*?[.。]',
            r'作为.*?，我需要.*?，以便.*?[。.]'
        ]
        
        for pattern in story_patterns:
            for match in re.finditer(pattern, content, re.IGNORECASE):
                user_stories.append({
                    'story': match.group().strip(),
                    'position': match.start()
                })
        
        return user_stories
